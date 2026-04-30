"""
generalroutines.py
------------------
Algemene hulpfuncties voor het aanmaken en opslaan van Word- en Excel-documenten.

Bevat functies om tekst in Word-documenten te vervangen (Replace), DataFrames als
tabel in een Word-document in te voegen (AddTable) en DataFrames op te slaan als
opgemaakte Excel-bestanden (SaveExcel).
"""

from docx import Document
from docx.shared import Inches


def Replace(replace_dict, doc):
    """Vervang tekst-sleutelwoorden in alle paragrafen van een Word-document.

    Args:
        replace_dict (dict): Woordenboek waarbij elke sleutel de te vervangen
            tekst is en de waarde de vervangende tekst.
        doc (docx.Document): Het te bewerken Word-document.

    Returns:
        docx.Document: Het bewerkte document met alle vervangingen toegepast.
    """
    for item in replace_dict:
        # Iterate through paragraphs and replace the text
        for paragraph in doc.paragraphs:
            if item in paragraph.text:
                paragraph.text = paragraph.text.replace(item, replace_dict[item])

    return doc

def AddTable(target_string, df, doc):
    """Voeg een DataFrame als opgemaakte tabel in een Word-document in.

    Zoekt de paragraaf die ``target_string`` bevat, plaatst daarna een tabel met
    de inhoud van ``df`` (kolomnamen als header) en verwijdert ``target_string``
    uit de paragraaf.  De eerste kolom krijgt een breedte van 4 inch, overige
    kolommen 2 inch.  Vereist de tabelstijl ``StijlHans`` in het document.

    Args:
        target_string (str): De tijdelijke plaatshouder-tekst in het document
            die de positie van de tabel aangeeft.
        df (pandas.DataFrame): Het DataFrame waarvan de inhoud als tabel wordt
            ingevoegd.
        doc (docx.Document): Het te bewerken Word-document.

    Returns:
        docx.Document: Het bewerkte document met de ingevoegde tabel.
    """
    # Find the paragraph that contains the target string
    target_paragraph = None

    for paragraph in doc.paragraphs:
        if target_string in paragraph.text:
            target_paragraph = paragraph
            break

    # add a table to the end and create a reference variable
    # extra row is so we can add the header row
    table = doc.add_table(df.shape[0]+1, df.shape[1],style='StijlHans')

    # Insert the table after the target paragraph
    p = target_paragraph._element
    new_tbl = table._tbl
    p.addnext(new_tbl)

    # Remove the target text if needed
    run = target_paragraph.runs[0]
    run.text = run.text.replace(target_string, '')


    # add the header rows.
    for j in range(df.shape[-1]):
        table.cell(0,j).text = df.columns[j]

    # add the rest of the data frame
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            table.cell(i+1,j).text = str(df.values[i,j])

    # Apply column widths
    i=0
    for cell in table.columns[0].cells:
        if i==0:
            cell.width = Inches(4)
        else:
            cell.width = Inches(2)
    return doc

def SaveExcel(df, filename, sheetname):
    """Sla een DataFrame op als opgemaakt Excel-bestand met automatische kolombreedtes.

    Maakt gebruik van XlsxWriter als engine.  Elke kolom krijgt een breedte
    gebaseerd op de maximale tekstlengte (begrensd op 150 tekens) met
    tekstterugloop ingeschakeld.

    Args:
        df (pandas.DataFrame): Het op te slaan DataFrame.
        filename (str): Pad naar het doelbestand (wordt aangemaakt of
            overschreven).
        sheetname (str): Naam van het werkblad in het Excel-bestand.
    """
    import pandas as pd

    # Create a Pandas Excel writer using XlsxWriter as the engine.
    writer = pd.ExcelWriter(filename, engine='xlsxwriter')

    # Convert the dataframe to an XlsxWriter Excel object.
    df.to_excel(writer, sheet_name=sheetname, index=False, na_rep='')

    # Get the xlsxwriter workbook and worksheet objects.
    workbook  = writer.book
    worksheet = writer.sheets[sheetname]

    # Add a format to wrap text.
    wrap_format = workbook.add_format({'text_wrap': True})

    for column in df:
        column_length = max(df[column].astype(str).map(len).max(), len(column))*1.02
        column_length = min(column_length, 150)
        col_idx = df.columns.get_loc(column)

        # Apply the format to the column.
        worksheet.set_column(col_idx, col_idx, column_length, wrap_format)

    # Close the Pandas Excel writer and output the Excel file.
    writer.close()
